import docx
import re
from docxtpl import DocxTemplate


class PasrsDoc(object):

    # 需要替换的内容:KEY,type
    # type image:0 text:1
    doc_parameters = list()


    # 开始符， 结束符 获取中间key
    def parse_parameters(self, param, start_with, end_with):
        type = 1
        if param == "":
            return

        match_str = "%s(.+?)%s" % (start_with, end_with)
        value = re.findall(match_str, param)
        if value is None:
            return "", -1

        for item in value:
            if item.startswith("image") or item.startswith("Image"):
                type = 0

            if {"value": item, "type": type} not in self.doc_parameters:
                self.doc_parameters.append({"value": item, "type": type})

    def parse_docx(self, filename, start_with, end_with):
        # found = False
        table_list = list()
        self.doc_parameters.clear()
        doc = docx.Document(filename)

        #普通文本处理
        for p in doc.paragraphs:
            self.parse_parameters(p.text, start_with, end_with)
        #表格处理
        for t in doc.tables:
            # if found:
            #     table_list.clear()
            for r in t.rows:
                temp = list()
                for c in r.cells:
                    if c.text not in temp:
                        temp.append({"text": c.text, "width": c.width})
                        self.parse_parameters(c.text, start_with, end_with)

                # if found:
                #     table_list.append(temp)

            # if found:
            #     found = False

        return table_list, self.doc_parameters


#dome

utils = PasrsDoc()
doc_parameters = utils.parse_docx('./templates/dome.docx','{{','}}')
datax = {
    'name': 'elk',
    'age': '18',
    '测试': 'aaa',
}
print("------------------------替换参数----------------------")
print(doc_parameters)

tpl = DocxTemplate('./templates/dome.docx')

tpl.render(datax)
tpl.save('./templates/newdome.docx')

